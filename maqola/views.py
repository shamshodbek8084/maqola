from django.shortcuts import render
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from .models import Maqola
from .serializers import Maqola_Serializer
from django.http import HttpResponse
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


# API orqali ko‘rish va qo‘shish
class MaqolaListCreateAPIView(APIView):
    def get(self, request):
        maqolalar = Maqola.objects.all().order_by('number')
        serializer = Maqola_Serializer(maqolalar, many=True)
        return Response(serializer.data)

    def post(self, request):
        serializer = Maqola_Serializer(data=request.data)
        if serializer.is_valid():
            serializer.save() 
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


# Word faylga dizaynli chiqarish
class ExportMaqolaDocxView(APIView):
    def get(self, request):
        maqolalar = Maqola.objects.all().order_by('number')
        doc = Document()
        last = maqolalar.last()

        heading_text = (
            "Muhammad al-Xorazmiy nomidagi Toshkent axborot texnologiyalari universiteti Kiberxavfsizlik fakulteti\n"
            f"{last.fakultet_raqami}"" - " f"“{last.fakultet} (Axborot-kommunikatsiya texnologiyalari va servis)” ta’lim yo‘nalishi  {last.guruh_raqami}-guruh "
            f"talabasi {last.talaba_fish}ning\n"
            "ILMIY ISHLARI RO‘YXATI"
        )
        heading = doc.add_paragraph(heading_text)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in heading.runs:
            run.bold = True
            run.font.size = Pt(12)

        # Jadval yaratish
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'

        # Sarlavhalar
        headers = [
            "№",
            "Ilmiy ishning nomi",
            "Bosma, qo‘l yozma yoki elektron",
            "Jurnal, to‘plam (yil, nomer, betlari), nashriyot yoki mualliflik guvohnomasi raqami",
            "Bosma taboq yoki betlar soni, mualliflik ishtiroki",
            "Mualliflarning F.I.Sh"
        ]
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            p = hdr_cells[i].paragraphs[0]
            run = p.add_run(header)
            run.bold = True
            run.font.size = Pt(10)

        # Har bir maqolani qo‘shish
        for m in maqolalar:
            row_cells = table.add_row().cells
            row_cells[0].text = str(m.number)
            row_cells[1].text = m.title
            row_cells[2].text = m.format
            row_cells[3].text = m.journal_information
            row_cells[4].text = m.piece
            row_cells[5].text = m.authors

        # HTTP Response qilib yuborish
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename=ilmiy_ishlar_ro‘yxati.docx'
        doc.save(response)
        return response

